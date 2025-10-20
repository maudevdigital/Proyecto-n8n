#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para convertir INFORME_COMPLETO_HU001.txt a formato DOCX
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_heading(doc, text, level=1):
    """Agregar encabezado con formato"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_with_format(doc, text, bold=False, italic=False, color=None, font_size=11):
    """Agregar párrafo con formato personalizado"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return paragraph

def is_section_header(line):
    """Detectar si una línea es un encabezado de sección"""
    # Detectar líneas como "1. INFORMACIÓN DEL PROYECTO"
    return re.match(r'^\d+\.\s+[A-ZÁÉÍÓÚÑ\s]+$', line.strip())

def is_subsection_header(line):
    """Detectar si una línea es un sub-encabezado"""
    # Detectar líneas como "3.1. Plataforma de Automatización"
    return re.match(r'^\d+\.\d+\.\s+.+$', line.strip())

def is_separator_line(line):
    """Detectar líneas separadoras (━━━━━)"""
    return '━' in line or '═' in line or '─' in line

def is_bullet_point(line):
    """Detectar puntos de lista"""
    stripped = line.strip()
    return stripped.startswith('•') or stripped.startswith('✓') or stripped.startswith('✗') or stripped.startswith('□')

def main():
    print("📄 Generando documento DOCX...")
    
    # Leer el archivo TXT
    with open('INFORME_COMPLETO_HU001.txt', 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Crear documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Variables de estado
    in_code_block = False
    skip_next = False
    
    for i, line in enumerate(lines):
        if skip_next:
            skip_next = False
            continue
            
        # Saltar líneas de diseño ASCII
        if line.strip().startswith('╔') or line.strip().startswith('║') or line.strip().startswith('╚'):
            continue
        
        # Detectar separadores
        if is_separator_line(line):
            # Agregar línea horizontal
            doc.add_paragraph('_' * 80)
            continue
        
        # Línea vacía
        if not line.strip():
            continue
        
        # Detectar encabezados de sección principal
        if is_section_header(line):
            # Saltar la línea separadora después si existe
            if i + 1 < len(lines) and is_separator_line(lines[i + 1]):
                skip_next = True
            add_heading(doc, line.strip(), level=1)
            continue
        
        # Detectar sub-encabezados
        if is_subsection_header(line):
            # Saltar la línea separadora después si existe
            if i + 1 < len(lines) and is_separator_line(lines[i + 1]):
                skip_next = True
            add_heading(doc, line.strip(), level=2)
            continue
        
        # Detectar listas con bullets
        if is_bullet_point(line):
            stripped = line.strip()
            # Determinar el símbolo y color
            if stripped.startswith('✓'):
                color = (0, 128, 0)  # Verde
                text = stripped[1:].strip()
            elif stripped.startswith('✗'):
                color = (255, 0, 0)  # Rojo
                text = stripped[1:].strip()
            elif stripped.startswith('□'):
                color = (128, 128, 128)  # Gris
                text = stripped[1:].strip()
            else:
                color = None
                text = stripped[1:].strip()
            
            paragraph = doc.add_paragraph(text, style='List Bullet')
            if color:
                paragraph.runs[0].font.color.rgb = RGBColor(*color)
            continue
        
        # Detectar bloques de código (líneas con indentación)
        if line.startswith('   ') or line.startswith('\t'):
            # Bloque de código
            paragraph = doc.add_paragraph(line.strip())
            paragraph.style = 'No Spacing'
            run = paragraph.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue
        
        # Texto normal
        text = line.strip()
        if text:
            # Detectar si es texto en negrita (líneas cortas en mayúsculas)
            if len(text) < 80 and text.isupper() and not is_separator_line(text):
                add_paragraph_with_format(doc, text, bold=True, font_size=12)
            else:
                doc.add_paragraph(text)
    
    # Guardar documento
    output_file = 'INFORME_COMPLETO_HU001.docx'
    doc.save(output_file)
    
    print(f"✅ Documento generado: {output_file}")
    print(f"📊 Tamaño: {len(lines)} líneas procesadas")

if __name__ == '__main__':
    main()
